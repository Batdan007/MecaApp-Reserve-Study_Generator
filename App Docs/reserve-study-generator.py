import os
import json
import datetime
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

class ReserveStudyGenerator:
    """
    A class to generate comprehensive reserve study reports for engineering companies,
    similar to MECA Engineering's SIRS reports.
    """
    def __init__(self):
        """Initialize the report generator with default settings."""
        self.property_data = {}
        self.component_data = []
        self.financial_data = {}
        self.template_dir = "templates/"
        self.output_dir = "output/"
        self.current_date = datetime.datetime.now()
        
        # Ensure output directory exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def load_property_data(self, file_path=None, data=None):
        """
        Load property data from a JSON file or directly from a dictionary.
        
        Args:
            file_path (str, optional): Path to the JSON file with property data.
            data (dict, optional): Property data as a dictionary.
        """
        if file_path:
            with open(file_path, 'r') as file:
                self.property_data = json.load(file)
        elif data:
            self.property_data = data
        else:
            raise ValueError("Either file_path or data must be provided")
    
    def load_component_data(self, file_path=None, data=None):
        """
        Load component inventory data from a JSON file or directly from a list.
        
        Args:
            file_path (str, optional): Path to the JSON file with component data.
            data (list, optional): Component data as a list of dictionaries.
        """
        if file_path:
            with open(file_path, 'r') as file:
                self.component_data = json.load(file)
        elif data:
            self.component_data = data
        else:
            raise ValueError("Either file_path or data must be provided")
    
    def load_financial_data(self, file_path=None, data=None):
        """
        Load financial data from a JSON file or directly from a dictionary.
        
        Args:
            file_path (str, optional): Path to the JSON file with financial data.
            data (dict, optional): Financial data as a dictionary.
        """
        if file_path:
            with open(file_path, 'r') as file:
                self.financial_data = json.load(file)
        elif data:
            self.financial_data = data
        else:
            raise ValueError("Either file_path or data must be provided")
    
    def calculate_fully_funded_balance(self):
        """
        Calculate the fully funded balance based on component effective age and replacement cost.
        
        Returns:
            float: The fully funded balance.
        """
        ffb = 0
        for component in self.component_data:
            useful_life = component.get("useful_life", 0)
            remaining_life = component.get("remaining_life", 0)
            replacement_cost = component.get("replacement_cost", 0)
            
            if useful_life > 0:
                effective_age = useful_life - remaining_life
                ffb += (effective_age / useful_life) * replacement_cost
        
        return ffb
    
    def calculate_percent_funded(self):
        """
        Calculate the percent funded based on current reserve balance and fully funded balance.
        
        Returns:
            float: The percent funded (0-100).
        """
        ffb = self.calculate_fully_funded_balance()
        current_balance = self.financial_data.get("starting_reserve_balance", 0)
        
        if ffb > 0:
            return (current_balance / ffb) * 100
        return 0
    
    def generate_expenditure_forecast(self, years=30):
        """
        Generate a forecast of expenditures for the specified number of years.
        
        Args:
            years (int): Number of years to forecast.
            
        Returns:
            pandas.DataFrame: DataFrame with annual expenditures by component.
        """
        # Create a DataFrame with years as columns
        current_year = self.current_date.year
        years_range = range(current_year, current_year + years)
        expenditure_df = pd.DataFrame(index=[comp["name"] for comp in self.component_data], 
                                      columns=years_range)
        expenditure_df = expenditure_df.fillna(0)
        
        # Fill in expenditure data
        for component in self.component_data:
            name = component["name"]
            remaining_life = component["remaining_life"]
            useful_life = component["useful_life"]
            replacement_cost = component["replacement_cost"]
            
            # Add initial replacement
            replacement_year = current_year + remaining_life
            if replacement_year in years_range:
                expenditure_df.at[name, replacement_year] = replacement_cost
            
            # Add future replacements
            while replacement_year + useful_life < current_year + years:
                replacement_year += useful_life
                if replacement_year in years_range:
                    expenditure_df.at[name, replacement_year] = replacement_cost
        
        return expenditure_df
    
    def calculate_funding_scenarios(self, years=30):
        """
        Calculate different funding scenarios (baseline, threshold, and fully funded).
        
        Args:
            years (int): Number of years to forecast.
            
        Returns:
            dict: Dictionary with funding scenarios data.
        """
        expenditure_df = self.generate_expenditure_forecast(years)
        total_annual_expenditures = expenditure_df.sum()
        
        # Get financial data
        starting_balance = self.financial_data.get("starting_reserve_balance", 0)
        current_contribution = self.financial_data.get("annual_contribution", 0)
        
        # Calculate fully funded balance
        ffb = self.calculate_fully_funded_balance()
        
        # Calculate baseline funding (reserve balance never below zero)
        baseline_balance = starting_balance
        baseline_contributions = []
        min_contribution = 0
        
        # Find minimum contribution to never go below zero
        for contribution in range(0, 200000, 1000):  # Test from 0 to 200k in 1k increments
            test_balance = starting_balance
            all_positive = True
            
            for year in range(years):
                current_year = self.current_date.year + year
                if current_year in total_annual_expenditures.index:
                    test_balance = test_balance + contribution - total_annual_expenditures[current_year]
                else:
                    test_balance += contribution
                
                if test_balance < 0:
                    all_positive = False
                    break
            
            if all_positive:
                min_contribution = contribution
                break
        
        # Calculate baseline scenario
        baseline_funding = {
            "annual_contribution": min_contribution,
            "balances": [starting_balance]
        }
        
        baseline_balance = starting_balance
        for year in range(years):
            current_year = self.current_date.year + year
            if current_year in total_annual_expenditures.index:
                baseline_balance = baseline_balance + min_contribution - total_annual_expenditures[current_year]
            else:
                baseline_balance += min_contribution
            
            baseline_funding["balances"].append(baseline_balance)
        
        # Calculate threshold funding (maintain 50% funded)
        threshold_contribution = min_contribution * 1.25  # 25% more than baseline
        threshold_funding = {
            "annual_contribution": threshold_contribution,
            "balances": [starting_balance]
        }
        
        threshold_balance = starting_balance
        for year in range(years):
            current_year = self.current_date.year + year
            if current_year in total_annual_expenditures.index:
                threshold_balance = threshold_balance + threshold_contribution - total_annual_expenditures[current_year]
            else:
                threshold_balance += threshold_contribution
            
            threshold_funding["balances"].append(threshold_balance)
        
        # Calculate fully funded scenario (maintain close to 100% funded)
        fully_funded_contribution = min_contribution * 1.35  # 35% more than baseline
        fully_funded_funding = {
            "annual_contribution": fully_funded_contribution,
            "balances": [starting_balance]
        }
        
        fully_funded_balance = starting_balance
        for year in range(years):
            current_year = self.current_date.year + year
            if current_year in total_annual_expenditures.index:
                fully_funded_balance = fully_funded_balance + fully_funded_contribution - total_annual_expenditures[current_year]
            else:
                fully_funded_balance += fully_funded_contribution
            
            fully_funded_funding["balances"].append(fully_funded_balance)
        
        return {
            "baseline": baseline_funding,
            "threshold": threshold_funding,
            "fully_funded": fully_funded_funding,
            "expenditures": total_annual_expenditures.to_dict(),
            "percent_funded": self.calculate_percent_funded()
        }
    
    def generate_funding_chart(self, funding_scenarios, filename="funding_chart.png"):
        """
        Generate a chart showing different funding scenarios over time.
        
        Args:
            funding_scenarios (dict): Funding scenarios data from calculate_funding_scenarios.
            filename (str): Output filename for the chart.
            
        Returns:
            str: Path to the generated chart file.
        """
        years = len(funding_scenarios["baseline"]["balances"])
        x = range(self.current_date.year, self.current_date.year + years)
        x = x[:len(funding_scenarios["baseline"]["balances"])]  # Ensure length matches
        
        plt.figure(figsize=(12, 6))
        
        # Plot expenditures as negative bars
        expenditures = [funding_scenarios["expenditures"].get(year, 0) for year in x[:-1]]
        expenditures = [-exp for exp in expenditures]  # Make them negative for the chart
        plt.bar(x[:-1], expenditures, color='red', alpha=0.7, label='Expenditures')
        
        # Plot funding scenarios
        plt.plot(x, funding_scenarios["baseline"]["balances"], label='Baseline Funding', color='blue')
        plt.plot(x, funding_scenarios["threshold"]["balances"], label='Threshold Funding', color='green')
        plt.plot(x, funding_scenarios["fully_funded"]["balances"], label='Fully Funded', color='gold')
        
        plt.axhline(y=0, color='black', linestyle='-', alpha=0.3)
        plt.title('Reserve Fund Projections')
        plt.xlabel('Year')
        plt.ylabel('Reserve Balance ($)')
        plt.legend()
        plt.grid(True, alpha=0.3)
        
        output_path = os.path.join(self.output_dir, filename)
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return output_path
    
    def generate_annual_expenditure_chart(self, filename="annual_expenditures.png"):
        """
        Generate a chart showing annual expenditures by component.
        
        Args:
            filename (str): Output filename for the chart.
            
        Returns:
            str: Path to the generated chart file.
        """
        expenditure_df = self.generate_expenditure_forecast()
        
        # Get annual expenditures
        annual_expenditures = expenditure_df.sum()
        
        plt.figure(figsize=(12, 6))
        bars = plt.bar(annual_expenditures.index, annual_expenditures.values, color='blue', alpha=0.7)
        
        # Add component breakdown for years with expenditures
        colors = plt.cm.tab10.colors
        bottom = np.zeros(len(annual_expenditures))
        
        for i, component in enumerate(expenditure_df.index):
            component_values = expenditure_df.loc[component].values
            if any(component_values > 0):  # Check if the component has any expenditures
                plt.bar(annual_expenditures.index, component_values, bottom=bottom, 
                      color=colors[i % len(colors)], label=component)
                bottom += component_values
        
        plt.title('Annual Expenditures')
        plt.xlabel('Year')
        plt.ylabel('Expenditure ($)')
        plt.legend(loc='upper right', bbox_to_anchor=(1.15, 1))
        plt.grid(True, alpha=0.3)
        
        # Rotate x-axis labels for better readability
        plt.xticks(rotation=45)
        
        output_path = os.path.join(self.output_dir, filename)
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return output_path
    
    def set_section_margins(self, section, top=1.0, bottom=1.0, left=1.0, right=1.0):
        """Set margins for a document section."""
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
    
    def generate_word_report(self, output_filename="reserve_study_report.docx"):
        """
        Generate a complete Word report with all sections.
        
        Args:
            output_filename (str): Filename for the output report.
            
        Returns:
            str: Path to the generated report.
        """
        doc = Document()
        
        # Set up the document
        section = doc.sections[0]
        self.set_section_margins(section)
        
        # Add title page
        self._add_title_page(doc)
        doc.add_page_break()
        
        # Add table of contents
        self._add_table_of_contents(doc)
        doc.add_page_break()
        
        # Add definitions
        self._add_definitions_section(doc)
        doc.add_page_break()
        
        # Add project overview
        self._add_project_overview(doc)
        doc.add_page_break()
        
        # Add component inventory
        self._add_component_inventory(doc)
        doc.add_page_break()
        
        # Add component assessment
        self._add_component_assessment(doc)
        doc.add_page_break()
        
        # Generate funding charts
        funding_scenarios = self.calculate_funding_scenarios()
        funding_chart_path = self.generate_funding_chart(funding_scenarios)
        expenditure_chart_path = self.generate_annual_expenditure_chart()
        
        # Add expenditures section
        self._add_expenditures_section(doc, expenditure_chart_path)
        doc.add_page_break()
        
        # Add reserve fund analysis
        self._add_reserve_fund_analysis(doc, funding_scenarios, funding_chart_path)
        doc.add_page_break()
        
        # Add future guidelines and conclusions
        self._add_future_guidelines(doc)
        doc.add_page_break()
        
        # Add disclosures
        self._add_disclosures(doc)
        
        # Save the document
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def _add_title_page(self, doc):
        """Add the title page to the document."""
        # Add logo placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Placeholder for logo (would be replaced with an actual logo image)
        run.add_text("[LOGO PLACEHOLDER]")
        
        # Add property name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.property_data.get("name", "Property Name"))
        run.font.size = Pt(32)
        run.font.bold = True
        
        # Add property address
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.property_data.get("address", "Property Address"))
        run.font.size = Pt(20)
        
        # Add "Structural Integrity Reserve Study"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Structural Integrity Reserve Study")
        run.font.size = Pt(24)
        run.font.bold = True
        
        # Add report information
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_info = f"MECA Report # {self.property_data.get('reference_id', 'REPORT-ID')}\n"
        report_info += f"Inspection Date: {self.property_data.get('inspection_date', self.current_date.strftime('%m/%d/%Y'))}\n"
        report_info += f"Report Issued: {self.current_date.strftime('%m/%d/%Y')}"
        run = p.add_run(report_info)
        run.font.size = Pt(12)
        
        # Add company footer information
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = "MECA Engineering Corporation\n"
        footer_text += "5237 Summerlin Commons Blvd\n"
        footer_text += "Suite 460 • Fort Myers, FL\n"
        footer_text += "T: (239) 275-2598"
        run = p.add_run(footer_text)
        run.font.size = Pt(12)
    
    def _add_table_of_contents(self, doc):
        """Add table of contents to the document."""
        doc.add_heading("Table of Contents", level=1)
        
        # Add formatted TOC entries
        toc_sections = [
            ("Definitions", 3),
            ("Project Overview and Compliance", 5),
            ("Overview of Physical Inspection", 7),
            ("Component Inventory", 8),
            ("Component Assessment", 10),
            ("Expenditures", 18),
            ("Reserve Fund Analysis", 20),
            ("Future Guidelines", 28),
            ("Disclosures", 29)
        ]
        
        # Create a table for the TOC
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        table.style = 'Table Grid'
        table.columns[0].width = Inches(5)
        table.columns[1].width = Inches(1)
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Description"
        header_cells[1].text = "Page"
        
        # Add TOC entries
        for section, page in toc_sections:
            row_cells = table.add_row().cells
            row_cells[0].text = section
            row_cells[1].text = str(page)
    
    def _add_definitions_section(self, doc):
        """Add the definitions section to the document."""
        doc.add_heading("Definitions", level=1)
        
        definitions = [
            ("Baseline Funding", "Establishing a reserve funding goal of allowing the reserve cash balance to approach but never fall below zero during the cash flow projection."),
            ("Effective Age", "The difference Useful Life minus the Remaining Useful Life; used to compute fully-funded balance; fraction of life \"used\" up for a given component."),
            ("Full Funding", "Setting a reserve funding goal to attain and maintain reserves at or near 100 percent funded."),
            ("Fully Funded Balance (FFB)", "The sum of Effective Age multiplied by current cost for all components. The Fully Funded Balance is directly proportional to the fraction of life \"used up\" for a given component."),
            ("Long-Life Components", "Reserve fund components with an estimated remaining useful life of more than 30 years from the date of the study being prepared."),
            ("Percent Funded", "The ratio of the actual or projected Reserve Fund Balance versus the Fully Funded Balance."),
            ("Remaining Useful Life (RUL)", "An estimate of the number of years a reserve component has until it becomes unusable or requires replacement."),
            ("Reserve Fund Balance", "Actual or projected funds that the association has identified for use to pay for the future repair or replacement of the components which the association is obligated to maintain."),
            ("Structural Integrity Reserve Study (SIRS)", "A study of the reserve funds required for future major repairs and replacement of the condominium property performed as required under s. 718.112(2)(g)."),
            ("Threshold Funding", "Establishing a reserve funding goal of keeping the reserve balance above a specified dollar or percent funded amount."),
            ("Useful Life (UL)", "An estimate of the number of years a reserve component will remain in service or serve its intended function.")
        ]
        
        for term, definition in definitions:
            p = doc.add_paragraph()
            run = p.add_run(f"{term} – ")
            run.font.bold = True
            p.add_run(definition)
            doc.add_paragraph()  # Add spacing between definitions
    
    def _add_project_overview(self, doc):
        """Add the project overview section to the document."""
        doc.add_heading("Project Overview and Compliance", level=1)
        
        # Project Introduction
        doc.add_heading("Project Introduction", level=2)
        intro_text = f"MECA Engineering performed a Structural Integrity Reserve Study (SIRS) for {self.property_data.get('name', 'the property')}, "
        intro_text += f"which consists of {self.property_data.get('building_count', 'multiple')} condominium buildings. "
        intro_text += f"The buildings are located at {self.property_data.get('address', 'the property address')}. "
        intro_text += f"A site visit and physical inspection of the common components of the property that are the subject of this SIRS were performed on {self.property_data.get('inspection_date', self.current_date.strftime('%B %d, %Y'))}. "
        intro_text += f"The condominium building was constructed in {self.property_data.get('construction_date', 'YEAR')}. "
        
        doc.add_paragraph(intro_text)
        
        # Frequency and Timeline
        doc.add_heading("Frequency and Timeline", level=2)
        freq_text = "Initial Study: This study has been conducted as required by Florida law and completed on "
        freq_text += f"{self.current_date.strftime('%B %d, %Y')}.\n\n"
        freq_text += "Subsequent Studies: The next structural integrity reserve study is due in 10 years from the Initial Study. "
        freq_text += "Thereafter, studies must be conducted every 10 years. Due to the association's projected contributions and the "
        freq_text += "projected expenditures over the next 12 years, MECA recommends updating this structural integrity study in 3 years."
        
        doc.add_paragraph(freq_text)
        
        # Personnel
        doc.add_heading("Personnel", level=2)
        personnel_text = "This Structural Integrity Reserve Study has been conducted in accordance with the requirements of Florida Statute Chapter 718. "
        personnel_text += "The study was performed by a qualified staff member, with the visual inspection of the condominium property conducted by an engineer licensed under Chapter 471."
        
        doc.add_paragraph(personnel_text)
        
        # Property Data
        doc.add_heading("Property Data", level=2)
        
        # Create a table for property data
        table = doc.add_table(rows=16, cols=2)
        table.style = 'Table Grid'
        
        # Fill the table with property information
        property_fields = [
            ("Property Name:", self.property_data.get("name", "")),
            ("Property Address:", self.property_data.get("address", "")),
            ("Type of Facility:", self.property_data.get("facility_type", "Multifamily residential condominiums")),
            ("Construction Date(s):", self.property_data.get("construction_date", "")),
            ("Number of Buildings:", self.property_data.get("building_count", "")),
            ("Number of Stories:", self.property_data.get("stories", "")),
            ("Number of Units:", self.property_data.get("unit_count", "")),
            ("Building(s) Area:", self.property_data.get("building_area", "")),
            ("Superstructure:", self.property_data.get("superstructure", "")),
            ("Roofing System:", self.property_data.get("roofing_system", "")),
            ("Exterior Façade:", self.property_data.get("exterior_facade", "")),
            ("Heating:", self.property_data.get("heating", "")),
            ("Cooling:", self.property_data.get("cooling", "")),
            ("Electrical Wiring:", self.property_data.get("electrical", "")),
            ("Fire Suppression:", self.property_data.get("fire_suppression", "N/A")),
            ("Date of Site Visit:", self.property_data.get("inspection_date", self.current_date.strftime("%m/%d/%Y")))
        ]
        
        for i, (label, value) in enumerate(property_fields):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
    
    def _add_component_inventory(self, doc):
        """Add the component inventory section to the document."""
        doc.add_heading("Component Inventory", level=1)
        
        # Add explanation text
        inventory_text = "In accordance with Florida Statute 718, the following common components were visually inspected considered as part of this structural integrity reserve study:\n\n"
        inventory_text += "a. Roof\n"
        inventory_text += "b. Structure, including load-bearing walls and other primary structural members and primary structural systems\n"
        inventory_text += "c. Fireproofing and fire protection systems\n"
        inventory_text += "d. Plumbing\n"
        inventory_text += "e. Electrical systems\n"
        inventory_text += "f. Waterproofing and exterior painting\n"
        inventory_text += "g. Exterior windows and doors\n"
        inventory_text += "h. Any other item that has a deferred maintenance expense or replacement cost that exceeds $10,000 and the failure to replace or maintain such item negatively affects the items listed in sub-subparagraphs a.-f., as determined by the visual inspection portion of the structural integrity reserve study\n\n"
        
        doc.add_paragraph(inventory_text)
        
        # Add additional methodology text
        methodology_text = "These components, their useful life and their remaining useful life were analyzed based upon the visually observed conditions. "
        methodology_text += "Cost estimates are based upon similarly scoped projects and market supported unit costs (e.g. RSM Means). "
        methodology_text += "All costs are based on present-dollar value. Component life valuations are based on market standards for the subject's property type and location and visually observed condition of each component in compliance with Florida Statutes. "
        methodology_text += "Long-life components with a remaining useful life beyond 30 years were not considered for this reserve study. "
        methodology_text += "Additionally, components which were reported to be the responsibility of the individual unit owners were not included. "
        methodology_text += "The following, Table A, is the component inventory for this structural integrity reserve study:"
        
        doc.add_paragraph(methodology_text)
        
        # Add component inventory table
        doc.add_heading("Table A - Component Inventory", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Component"
        header_cells[1].text = "Qty."
        header_cells[2].text = "Unit"
        header_cells[3].text = "UL\n(years)"
        header_cells[4].text = "RUL\n(years)"
        
        # Add component rows
        for component in self.component_data:
            row_cells = table.add_row().cells
            row_cells[0].text = component.get("name", "")
            row_cells[1].text = str(component.get("quantity", ""))
            row_cells[2].text = component.get("unit", "")
            row_cells[3].text = str(component.get("useful_life", ""))
            row_cells[4].text = str(component.get("remaining_life", ""))
        
        # Add notes
        doc.add_paragraph("Notes:\nUL – Useful Life\nRUL – Remaining Useful Life")
    
    def _add_component_assessment(self, doc):
        """Add the component assessment section to the document."""
        doc.add_heading("Component Assessment", level=1)
        
        # Add assessment for each component
        for component in self.component_data:
            doc.add_heading(f"Exterior Building Components: {component.get('name', 'Component')}", level=2)
            
            # Add component description
            description = component.get("description", "No description provided.")
            doc.add_paragraph(description)
            
            # Add component details in a simple format
            details_text = f"Useful Life:\n{component.get('useful_life', 'N/A')} years\n\n"
            details_text += f"Remaining Useful Life:\n{component.get('remaining_life', 'N/A')} years\n\n"
            details_text += f"Quantity:\n{component.get('quantity', 'N/A')} {component.get('unit', '')}\n\n"
            details_text += f"Estimated Cost:\n${component.get('replacement_cost', 0):,.2f}"
            
            doc.add_paragraph(details_text)
            
            # Add placeholder for component image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_text("[COMPONENT IMAGE PLACEHOLDER]")
            
            # Add page break between components
            if component != self.component_data[-1]:  # Don't add page break after the last component
                doc.add_page_break()
    
    def _add_expenditures_section(self, doc, chart_path):
        """Add the expenditures section to the document."""
        doc.add_heading("Expenditures", level=1)
        
        # Add explanation text
        expenditure_text = "This SIRS projects expenditures for the components identified in the component inventory over the next 30 years. "
        expenditure_text +=